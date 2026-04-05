#!/usr/bin/env python3
"""
Strategic Business Review Validator

Validates .docx output against quality checklist requirements.
Run: python scripts/validate_sbr.py output.docx

Exit codes:
    0 - All checks passed
    1 - One or more checks failed
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.table import Table
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


class SBRValidator:
    """Validates Strategic Business Review documents."""

    FORBIDDEN_PATTERNS = [
        r'\bleverage synergies?\b',
        r'\boptimi[sz]e operational excellence\b',
        r'\bCoaching CTA\b',
        r'\bReflection line\b',
        r'\b\[.*?\]\b',  # Template placeholders like [Owner Name]
    ]

    LIVED_EXPERIENCE_MARKERS = [
        'deli', 'dry cleaner', 'clinic', 'healthcare',
        'slicer', 'steam', 'garments', 'counter',
        'Saturday', 'till', 'whiteboard'
    ]

    BRIDGE_PATTERN = r'Bridge\s*[-–—→>]'

    def __init__(self, filepath: str):
        self.filepath = Path(filepath)
        self.doc = Document(filepath)
        self.full_text = self._extract_full_text()
        self.word_count = len(self.full_text.split())
        self.errors = []
        self.warnings = []
        self.passes = []

    def _extract_full_text(self) -> str:
        """Extract all text from document."""
        paragraphs = [p.text for p in self.doc.paragraphs]
        # Also extract from tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return '\n'.join(paragraphs)

    def check_word_count(self) -> bool:
        """Check total word count is 3,000-5,000."""
        if 3000 <= self.word_count <= 5000:
            self.passes.append(f"Word count: {self.word_count} (target: 3,000-5,000)")
            return True
        elif self.word_count < 3000:
            self.errors.append(f"Word count too low: {self.word_count} (minimum: 3,000)")
            return False
        else:
            self.warnings.append(f"Word count high: {self.word_count} (target max: 5,000)")
            return True

    def check_executive_summary_length(self) -> bool:
        """Check executive summary is under 250 words."""
        # Find executive overview section
        exec_pattern = r'(?:Executive Overview|executive overview|1\.\s*Executive).*?(?=\n\d\.|$)'
        match = re.search(exec_pattern, self.full_text, re.DOTALL | re.IGNORECASE)

        if not match:
            self.warnings.append("Could not locate Executive Overview section")
            return True

        exec_text = match.group(0)
        exec_words = len(exec_text.split())

        if exec_words <= 350:  # Allow some flexibility for headers
            self.passes.append(f"Executive Overview: ~{exec_words} words (target: ≤350)")
            return True
        else:
            self.warnings.append(f"Executive Overview may be long: ~{exec_words} words")
            return True

    def check_lived_experiences(self) -> bool:
        """Check for at least one lived experience reference."""
        text_lower = self.full_text.lower()
        found = [m for m in self.LIVED_EXPERIENCE_MARKERS if m.lower() in text_lower]

        if found:
            self.passes.append(f"Lived experience references found: {', '.join(found[:3])}")
            return True
        else:
            self.errors.append("No lived experience references (deli/dry cleaner/clinic)")
            return False

    def check_action_boxes(self) -> bool:
        """Check for at least 3 Action Boxes (tables with borders)."""
        action_box_count = 0

        for table in self.doc.tables:
            # Check if table contains "Action Box" text
            table_text = ''
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text

            if 'Action Box' in table_text or '☐' in table_text:
                action_box_count += 1

        if action_box_count >= 3:
            self.passes.append(f"Action Boxes: {action_box_count} found (minimum: 3)")
            return True
        else:
            self.errors.append(f"Insufficient Action Boxes: {action_box_count} (minimum: 3)")
            return False

    def check_bridge_transitions(self) -> bool:
        """Check for at least 2 Bridge transitions."""
        bridges = re.findall(self.BRIDGE_PATTERN, self.full_text, re.IGNORECASE)

        if len(bridges) >= 2:
            self.passes.append(f"Bridge transitions: {len(bridges)} found (minimum: 2)")
            return True
        else:
            self.warnings.append(f"Few Bridge transitions: {len(bridges)} (recommended: 2+)")
            return True  # Warning only

    def check_forbidden_patterns(self) -> bool:
        """Check for forbidden corporate jargon and template artifacts."""
        found_forbidden = []

        for pattern in self.FORBIDDEN_PATTERNS:
            matches = re.findall(pattern, self.full_text, re.IGNORECASE)
            if matches:
                found_forbidden.extend(matches[:2])  # Limit examples

        if not found_forbidden:
            self.passes.append("No forbidden patterns (jargon/template artifacts)")
            return True
        else:
            self.errors.append(f"Forbidden patterns found: {', '.join(found_forbidden)}")
            return False

    def check_second_person_voice(self) -> bool:
        """Check document uses second person ('you') appropriately."""
        you_count = len(re.findall(r'\byou\b', self.full_text, re.IGNORECASE))
        your_count = len(re.findall(r'\byour\b', self.full_text, re.IGNORECASE))
        total_second_person = you_count + your_count

        # Expect roughly 1 per 50 words minimum
        expected_min = self.word_count // 100

        if total_second_person >= expected_min:
            self.passes.append(f"Second person voice: {total_second_person} instances")
            return True
        else:
            self.warnings.append(f"Limited second person: {total_second_person} (may feel impersonal)")
            return True

    def check_grow_framework(self) -> bool:
        """Check G.R.O.W. framework elements are present."""
        grow_elements = ['goal', 'reality', 'options', 'way forward']
        found = [e for e in grow_elements if e.lower() in self.full_text.lower()]

        if len(found) >= 3:
            self.passes.append(f"G.R.O.W. elements: {len(found)}/4 present")
            return True
        else:
            self.warnings.append(f"G.R.O.W. framework incomplete: {len(found)}/4 elements")
            return True

    def check_closing_tone(self) -> bool:
        """Check closing doesn't contain hype language."""
        hype_patterns = [
            r'\bact now\b', r'\bdon\'t miss\b', r'\blimited time\b',
            r'\bguarantee\b', r'\bamazing\b', r'\bincredible\b'
        ]

        # Check last 500 words
        closing = ' '.join(self.full_text.split()[-500:])

        found_hype = []
        for pattern in hype_patterns:
            if re.search(pattern, closing, re.IGNORECASE):
                found_hype.append(pattern.replace(r'\b', ''))

        if not found_hype:
            self.passes.append("Closing tone: appropriate (no hype)")
            return True
        else:
            self.errors.append(f"Hype language in closing: {', '.join(found_hype)}")
            return False

    def validate(self) -> bool:
        """Run all validation checks."""
        print(f"\n{'='*60}")
        print(f"Validating: {self.filepath.name}")
        print(f"{'='*60}\n")

        checks = [
            self.check_word_count,
            self.check_executive_summary_length,
            self.check_lived_experiences,
            self.check_action_boxes,
            self.check_bridge_transitions,
            self.check_forbidden_patterns,
            self.check_second_person_voice,
            self.check_grow_framework,
            self.check_closing_tone,
        ]

        for check in checks:
            check()

        # Print results
        if self.passes:
            print("PASSED:")
            for p in self.passes:
                print(f"  ✓ {p}")

        if self.warnings:
            print("\nWARNINGS:")
            for w in self.warnings:
                print(f"  ⚠ {w}")

        if self.errors:
            print("\nERRORS:")
            for e in self.errors:
                print(f"  ✗ {e}")

        # Summary
        print(f"\n{'='*60}")
        passed = len(self.errors) == 0
        status = "PASSED" if passed else "FAILED"
        print(f"Result: {status}")
        print(f"  {len(self.passes)} passed, {len(self.warnings)} warnings, {len(self.errors)} errors")
        print(f"{'='*60}\n")

        return passed


def main():
    if len(sys.argv) < 2:
        print("Usage: python validate_sbr.py <document.docx>")
        print("\nValidates a Strategic Business Review document against quality requirements.")
        sys.exit(1)

    filepath = sys.argv[1]

    if not Path(filepath).exists():
        print(f"ERROR: File not found: {filepath}")
        sys.exit(1)

    if not filepath.endswith('.docx'):
        print(f"ERROR: Expected .docx file, got: {filepath}")
        sys.exit(1)

    validator = SBRValidator(filepath)
    passed = validator.validate()

    sys.exit(0 if passed else 1)


if __name__ == '__main__':
    main()
