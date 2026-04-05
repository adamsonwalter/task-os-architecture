#!/usr/bin/env python3
"""
Strategic Business Review Document Generator

Template script for generating .docx reports with proper formatting.
Provides helper functions for Action Boxes, styling, and structure.

Usage:
    from generate_docx import SBRDocument

    doc = SBRDocument("Client Name")
    doc.add_opening_letter("Owner Name", "personalised opening...")
    doc.add_section("Executive Overview", content, action_items=[...])
    doc.save("output.docx")
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from datetime import datetime
from typing import List, Optional
import json
import os


class SBRDocument:
    """Strategic Business Review document generator with Morgan Management styling."""

    def __init__(self, client_name: str, config_path: str = None):
        self.client_name = client_name
        self.doc = Document()
        self.config = self._load_config(config_path)
        self._setup_page()
        self._setup_styles()
        self._setup_headers_footers()

    def _load_config(self, config_path: str = None) -> dict:
        """Load configuration from document-settings.json."""
        if config_path is None:
            # Default path relative to script location
            script_dir = os.path.dirname(os.path.abspath(__file__))
            config_path = os.path.join(script_dir, '..', 'config', 'document-settings.json')
        
        try:
            with open(config_path, 'r') as f:
                return json.load(f)
        except FileNotFoundError:
            # Return defaults if config not found
            return {
                "word": {
                    "header": {
                        "text": "Strategic Business Review for {{client-call-name}}",
                        "firstPageDifferent": True
                    },
                    "footer": {
                        "left": "justine@morganmg.com.au",
                        "center": "{{date}}",
                        "right": "Page {{page}}",
                        "dateFormat": "dd-MMM-yyyy",
                        "allPagesSame": True
                    }
                }
            }

    def _setup_page(self):
        """Configure A4 page with standard margins."""
        for section in self.doc.sections:
            section.page_width = Cm(21.0)   # A4 width
            section.page_height = Cm(29.7)  # A4 height
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)

    def _setup_styles(self):
        """Configure document styles - Arial throughout."""
        # Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Heading 1
        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 51, 102)

        # Heading 2
        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 51, 102)

        # Heading 3
        h3 = self.doc.styles['Heading 3']
        h3.font.name = 'Arial'
        h3.font.size = Pt(12)
        h3.font.bold = True

    def _setup_headers_footers(self):
        """Configure headers and footers from config."""
        section = self.doc.sections[0]
        
        # First page different
        if self.config.get('word', {}).get('header', {}).get('firstPageDifferent', True):
            section.different_first_page_header_footer = True
        
        # Regular header
        header = section.header
        header_text = self.config.get('word', {}).get('header', {}).get('text', '')
        header_text = header_text.replace('{{client-call-name}}', self.client_name)
        
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = header_text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Footer
        footer = section.footer
        footer_config = self.config.get('word', {}).get('footer', {})
        
        # Create three-column footer using a table
        footer_table = footer.add_table(rows=1, cols=3, width=Cm(16))
        footer_table.autofit = True
        
        # Left cell
        left_cell = footer_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.text = footer_config.get('left', 'justine@morganmg.com.au')
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Center cell (date)
        center_cell = footer_table.cell(0, 1)
        center_para = center_cell.paragraphs[0]
        date_format = footer_config.get('dateFormat', 'dd-MMM-yyyy')
        # Convert Python strftime format
        py_format = date_format.replace('dd', '%d').replace('MMM', '%b').replace('yyyy', '%Y')
        center_para.text = datetime.now().strftime(py_format)
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Right cell (page number - will be a field)
        right_cell = footer_table.cell(0, 2)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._add_page_number(right_para)

    def _add_page_number(self, paragraph):
        """Add page number field to paragraph."""
        run = paragraph.add_run("Page ")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        
        # Add PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def add_title(self, subtitle: Optional[str] = None):
        """Add document title with optional subtitle."""
        title = self.doc.add_heading(f"Strategic Business Review", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if subtitle:
            p = self.doc.add_paragraph(subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add prepared for line
        date_str = datetime.now().strftime("%B %Y")
        prepared = self.doc.add_paragraph(f"Prepared for {self.client_name} — {date_str}")
        prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()  # Spacer

    def add_opening_letter(self, owner_name: str, content: str):
        """Add personalised opening letter."""
        # Salutation
        p = self.doc.add_paragraph()
        run = p.add_run(f"Dear {owner_name},")
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)

        self.doc.add_paragraph()

        # Content paragraphs
        for para in content.split('\n\n'):
            if para.strip():
                p = self.doc.add_paragraph(para.strip())
                p.style = 'Normal'

        self.doc.add_paragraph()

    def add_section(
        self,
        title: str,
        content: str,
        level: int = 1,
        action_items: Optional[List[str]] = None,
        action_title: Optional[str] = None,
        reflective_question: Optional[str] = None
    ):
        """
        Add a report section with optional Action Box.

        Args:
            title: Section heading
            content: Section body text (use \\n\\n for paragraph breaks)
            level: Heading level (1 or 2)
            action_items: List of checkbox items for Action Box
            action_title: Title for Action Box (defaults to section title theme)
            reflective_question: Optional reflection prompt to add after Action Box
        """
        # Heading
        self.doc.add_heading(title, level=level)

        # Content paragraphs
        for para in content.split('\n\n'):
            if para.strip():
                p = self.doc.add_paragraph(para.strip())

        # Action Box if items provided
        if action_items:
            box_title = action_title or title.split('—')[0].strip()
            self.add_action_box(box_title, action_items)

        # Reflective question
        if reflective_question:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run(f"Reflective call-out: ")
            run.italic = True
            run.font.name = 'Arial'
            p.add_run(reflective_question).font.name = 'Arial'

    def add_action_box(self, title: str, items: List[str]):
        """
        Add bordered Action Box that won't split across pages.

        Args:
            title: Box title (e.g., "Client Care Rhythm")
            items: List of checkbox items
        """
        self.doc.add_paragraph()  # Spacer

        # Create single-cell table
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = True

        # Set table borders
        self._set_table_borders(table)

        # Prevent row from splitting across pages
        self._prevent_row_split(table.rows[0])

        # Add content to cell
        cell = table.cell(0, 0)

        # Title paragraph
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run(f"Action Box — {title}")
        title_run.bold = True
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(11)

        # Checkbox items
        for item in items:
            para = cell.add_paragraph()
            run = para.add_run(f"☐ {item}")
            run.font.name = 'Arial'
            run.font.size = Pt(11)

        self.doc.add_paragraph()  # Spacer after

    def _set_table_borders(self, table):
        """Apply borders to table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)

        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _prevent_row_split(self, row):
        """Prevent table row from splitting across pages."""
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)

    def add_bridge(self, text: str):
        """Add a Bridge transition sentence."""
        p = self.doc.add_paragraph()
        run = p.add_run(f"Bridge → ")
        run.italic = True
        run.font.name = 'Arial'
        p.add_run(text).font.name = 'Arial'
        self.doc.add_paragraph()

    def add_quarter_block(
        self,
        quarter: str,
        months: str,
        focus: str,
        actions: List[str],
        outcome: str
    ):
        """Add a quarterly roadmap block."""
        self.doc.add_heading(f"{quarter} ({months})", level=2)

        # Focus
        p = self.doc.add_paragraph()
        run = p.add_run("Focus: ")
        run.bold = True
        p.add_run(focus)

        # Actions
        p = self.doc.add_paragraph()
        run = p.add_run("Actions:")
        run.bold = True

        for action in actions:
            p = self.doc.add_paragraph(f"• {action}")
            p.paragraph_format.left_indent = Cm(0.5)

        # Outcome
        p = self.doc.add_paragraph()
        run = p.add_run("Outcome: ")
        run.bold = True
        run.italic = True
        p.add_run(outcome).italic = True

        self.doc.add_paragraph()

    def add_final_cta(self, owner_name: str, message: str):
        """Add final call-to-action box."""
        self.doc.add_paragraph()

        table = self.doc.add_table(rows=1, cols=1)
        self._set_table_borders(table)
        self._prevent_row_split(table.rows[0])

        cell = table.cell(0, 0)

        # Title
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run("What's Next?")
        title_run.bold = True
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(12)

        # Message
        para = cell.add_paragraph()
        run = para.add_run(f"{owner_name}, {message}")
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    def add_closing_signature(self):
        """Add Dr. Justine Hicks signature block."""
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        p = self.doc.add_paragraph("With clarity and courage,")

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        run = p.add_run("Dr. Justine Hicks")
        run.bold = True
        run.font.name = 'Arial'

        p = self.doc.add_paragraph("Morgan Management Group")
        p.runs[0].font.name = 'Arial'

    def save(self, filepath: str):
        """Save document to file."""
        self.doc.save(filepath)
        print(f"Document saved: {filepath}")


# Example usage
if __name__ == '__main__':
    # Demonstrate structure
    doc = SBRDocument("Example Business Pty Ltd")

    doc.add_title("Morgan Management Group")

    doc.add_opening_letter(
        "Jane",
        """Thank you for taking the time to complete the Strategic Business Review survey. What follows is not a report to admire from a distance — it's a rhythm to live.

Your business has strong bones. Revenue is growing, team culture feels right, and you've built something real. But there's a quiet warning in your answers: the dashboard isn't readable yet, and cashflow feels like weather rather than a forecast.

The fix is rarely more hustle. It's rhythm, visibility, and a few small levers pulled consistently."""
    )

    doc.add_section(
        "1. Executive Overview — the three truths I want you to hold steady",
        """Your business is growing, but growth without rhythm creates fragility. The till rings, but margin discipline is thin.

Truth one: You have strong hands on the wheel, but limited dashboard visibility.

Truth two: Your team culture is genuine, but undocumented — which means it depends on your presence.

Truth three: Cashflow is your current constraint, not your weakness.

So the leverage points are clear: weekly financial rhythm, documented client care standards, and one "fast yes" offer to accelerate cash.""",
        action_items=[
            "Review weekly cash position every Friday (20 mins)",
            "Document one client touchpoint standard this week",
            "Define your 'fast yes' offer in one paragraph"
        ],
        action_title="Executive Focus"
    )

    doc.add_bridge("With these truths held gently but firmly, we can look at where you actually are in the lifecycle.")

    doc.save("example_sbr.docx")
    print("\nExample document created: example_sbr.docx")
